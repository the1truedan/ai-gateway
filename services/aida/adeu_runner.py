#!/usr/bin/env python3
"""Adeu DOCX redline integration for A.I.D.A. — Phase 3.

MIT: https://github.com/dealfluence/adeu
DOCX ↔ Markdown / Track Changes for style-normalized letters and caregiver drafts.

Uses local adeu CLI (venv) or SDK. Prepare-only; HITL before distribution.
"""

from __future__ import annotations

import json
import os
import shutil
import subprocess
import time
from pathlib import Path
from typing import Any

ADEU_DISABLE = os.environ.get("AIDA_ADEU_DISABLE", "0") == "1"
ADEU_AUTHOR = os.environ.get("AIDA_ADEU_AUTHOR", "A.I.D.A.").strip() or "A.I.D.A."
ADEU_TIMEOUT = float(os.environ.get("AIDA_ADEU_TIMEOUT", "120"))


def _adeu_bin() -> str | None:
    env = os.environ.get("AIDA_ADEU_CMD", "").strip()
    if env and Path(env).is_file():
        return env
    which = shutil.which("adeu")
    if which:
        return which
    # Prefer sibling venv next to this module
    here = Path(__file__).resolve().parent
    for cand in (
        here / ".venv" / "bin" / "adeu",
        here.parent.parent / "services" / "aida" / ".venv" / "bin" / "adeu",
    ):
        if cand.is_file():
            return str(cand)
    # python -m not always available; try import entry
    try:
        import adeu  # noqa: F401

        return "sdk"
    except ImportError:
        return None


def adeu_available() -> dict[str, Any]:
    if ADEU_DISABLE:
        return {
            "available": False,
            "reason": "AIDA_ADEU_DISABLE=1",
            "license": "MIT",
        }
    bin_path = _adeu_bin()
    if not bin_path:
        return {
            "available": False,
            "reason": "adeu not installed",
            "hint": "pip install adeu  # in services/aida/.venv",
            "license": "MIT",
        }
    version = None
    if bin_path != "sdk":
        try:
            r = subprocess.run(
                [bin_path, "-v"],
                capture_output=True,
                text=True,
                timeout=10,
            )
            version = (r.stdout or r.stderr or "").strip().splitlines()[:1]
            version = version[0] if version else None
        except Exception as exc:  # noqa: BLE001
            version = f"probe_error:{exc}"[:80]
    else:
        try:
            import adeu

            version = getattr(adeu, "__version__", "sdk")
        except Exception:  # noqa: BLE001
            version = "sdk"
    return {
        "available": True,
        "bin": bin_path if bin_path != "sdk" else None,
        "mode": "cli" if bin_path != "sdk" else "sdk",
        "version": version,
        "author_default": ADEU_AUTHOR,
        "license": "MIT",
        "ops": ["extract", "apply", "sanitize", "md_to_docx", "from_brief"],
        "note": (
            "DOCX Track Changes redline. Prepare-only; HITL before send. "
            "Local only — no Adeu Cloud."
        ),
    }


def extract_docx(
    docx_path: str | Path,
    *,
    out_path: str | Path | None = None,
    clean_view: bool = True,
) -> dict[str, Any]:
    """Extract Markdown/text from DOCX via adeu extract."""
    t0 = time.perf_counter()
    src = Path(docx_path).expanduser().resolve()
    if not src.is_file():
        return {"status": "error", "error": f"not found: {src}"}
    avail = adeu_available()
    if not avail.get("available"):
        return {"status": "unavailable", **{k: avail.get(k) for k in ("reason", "hint")}}

    dest = Path(out_path).expanduser().resolve() if out_path else src.with_suffix(".md")
    dest.parent.mkdir(parents=True, exist_ok=True)

    bin_path = _adeu_bin()
    if bin_path and bin_path != "sdk":
        cmd = [bin_path, "extract", str(src), "-o", str(dest)]
        if clean_view:
            cmd.append("--clean-view")
        try:
            r = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=ADEU_TIMEOUT,
            )
            if r.returncode != 0:
                return {
                    "status": "error",
                    "error": (r.stderr or r.stdout or "extract failed")[:500],
                    "returncode": r.returncode,
                }
            text = dest.read_text(encoding="utf-8") if dest.is_file() else ""
            return {
                "status": "ok",
                "engine": "adeu-cli",
                "source": str(src),
                "output": str(dest),
                "chars": len(text),
                "markdown_preview": text[:4000],
                "elapsed_ms": round((time.perf_counter() - t0) * 1000, 1),
            }
        except Exception as exc:  # noqa: BLE001
            return {"status": "error", "error": str(exc)[:400]}

    return {"status": "error", "error": "adeu CLI required for extract"}


def apply_edits(
    docx_path: str | Path,
    edits: list[dict[str, Any]] | Path | str,
    *,
    out_path: str | Path | None = None,
    author: str | None = None,
    dry_run: bool = False,
) -> dict[str, Any]:
    """Apply Track Changes edits. edits = list of dicts or path to JSON file.

    Edit shape (SDK/CLI compatible):
      {"type": "modify", "target_text": "...", "new_text": "...", "comment": "..."}
    or ModifyText fields: target_text, new_text, comment
    """
    t0 = time.perf_counter()
    src = Path(docx_path).expanduser().resolve()
    if not src.is_file():
        return {"status": "error", "error": f"not found: {src}"}
    avail = adeu_available()
    if not avail.get("available"):
        return {"status": "unavailable", **{k: avail.get(k) for k in ("reason", "hint")}}

    author = (author or ADEU_AUTHOR).strip() or ADEU_AUTHOR
    dest = (
        Path(out_path).expanduser().resolve()
        if out_path
        else src.with_name(f"{src.stem}.redlined.docx")
    )
    dest.parent.mkdir(parents=True, exist_ok=True)

    # Normalize edits to file
    edits_path: Path | None = None
    edit_list: list[dict[str, Any]] = []
    if isinstance(edits, (str, Path)):
        ep = Path(edits).expanduser().resolve()
        if not ep.is_file():
            return {"status": "error", "error": f"edits file not found: {ep}"}
        edits_path = ep
        try:
            raw = json.loads(ep.read_text(encoding="utf-8"))
            if isinstance(raw, list):
                edit_list = raw
            elif isinstance(raw, dict) and "edits" in raw:
                edit_list = list(raw["edits"])
            else:
                return {"status": "error", "error": "edits JSON must be list or {edits:[]}"}
        except json.JSONDecodeError as exc:
            return {"status": "error", "error": f"invalid edits JSON: {exc}"}
    else:
        edit_list = list(edits)
        edits_path = dest.parent / f"{src.stem}.edits.json"
        # Normalize type field
        norm: list[dict[str, Any]] = []
        for e in edit_list:
            if not isinstance(e, dict):
                continue
            item = dict(e)
            if "type" not in item and ("target_text" in item or "new_text" in item):
                item["type"] = "modify"
            norm.append(item)
        edit_list = norm
        edits_path.write_text(json.dumps(edit_list, indent=2), encoding="utf-8")

    # Prefer SDK for structured apply
    try:
        from adeu import ModifyText, RedlineEngine
        from io import BytesIO

        stream = BytesIO(src.read_bytes())
        engine = RedlineEngine(stream, author=author)
        objs = []
        for e in edit_list:
            et = (e.get("type") or "modify").lower()
            if et in ("modify", "modifytext", "replace"):
                kwargs: dict[str, Any] = {
                    "target_text": str(e.get("target_text") or e.get("old") or ""),
                    "new_text": str(e.get("new_text") or e.get("new") or ""),
                }
                if e.get("comment"):
                    kwargs["comment"] = str(e["comment"])
                if e.get("match_mode"):
                    kwargs["match_mode"] = e["match_mode"]
                objs.append(ModifyText(**kwargs))
            # other types left for CLI path
        if objs and not dry_run:
            engine.apply_edits(objs)
            dest.write_bytes(engine.save_to_stream().getvalue())
            return {
                "status": "ok",
                "engine": "adeu-sdk",
                "source": str(src),
                "output": str(dest),
                "edits_path": str(edits_path),
                "edits_applied": len(objs),
                "author": author,
                "dry_run": False,
                "hitl_required": True,
                "decision_authority": "prepare_only",
                "elapsed_ms": round((time.perf_counter() - t0) * 1000, 1),
            }
        if dry_run and objs:
            return {
                "status": "ok",
                "engine": "adeu-sdk",
                "dry_run": True,
                "edits_preview": edit_list[:50],
                "edits_count": len(edit_list),
                "author": author,
                "hitl_required": True,
                "elapsed_ms": round((time.perf_counter() - t0) * 1000, 1),
            }
    except Exception as sdk_exc:  # noqa: BLE001
        sdk_err = str(sdk_exc)[:300]
    else:
        sdk_err = None

    # CLI fallback
    bin_path = _adeu_bin()
    if not bin_path or bin_path == "sdk":
        return {
            "status": "error",
            "error": f"SDK apply failed: {sdk_err}; no CLI",
        }

    cmd = [
        bin_path,
        "apply",
        str(src),
        str(edits_path),
        "-o",
        str(dest),
        "--author",
        author,
    ]
    if dry_run:
        cmd.append("--dry-run")
    try:
        r = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=ADEU_TIMEOUT,
        )
        ok = r.returncode == 0 and (dry_run or dest.is_file())
        return {
            "status": "ok" if ok else "error",
            "engine": "adeu-cli",
            "source": str(src),
            "output": str(dest) if dest.is_file() else None,
            "edits_path": str(edits_path),
            "author": author,
            "dry_run": dry_run,
            "stdout": (r.stdout or "")[:1000],
            "stderr": (r.stderr or "")[:500],
            "returncode": r.returncode,
            "sdk_error": sdk_err,
            "hitl_required": True,
            "decision_authority": "prepare_only",
            "elapsed_ms": round((time.perf_counter() - t0) * 1000, 1),
        }
    except Exception as exc:  # noqa: BLE001
        return {"status": "error", "error": str(exc)[:400], "sdk_error": sdk_err}


def sanitize_docx(
    docx_path: str | Path,
    *,
    out_path: str | Path | None = None,
    author: str | None = None,
    keep_markup: bool = True,
    accept_all: bool = False,
) -> dict[str, Any]:
    """Strip metadata / prep for distribution (still HITL)."""
    t0 = time.perf_counter()
    src = Path(docx_path).expanduser().resolve()
    if not src.is_file():
        return {"status": "error", "error": f"not found: {src}"}
    avail = adeu_available()
    if not avail.get("available"):
        return {"status": "unavailable", **{k: avail.get(k) for k in ("reason", "hint")}}

    dest = (
        Path(out_path).expanduser().resolve()
        if out_path
        else src.with_name(f"{src.stem}.sanitized.docx")
    )
    dest.parent.mkdir(parents=True, exist_ok=True)
    bin_path = _adeu_bin()
    if not bin_path or bin_path == "sdk":
        return {"status": "error", "error": "adeu CLI required for sanitize"}

    cmd = [bin_path, "sanitize", str(src), "-o", str(dest)]
    if keep_markup:
        cmd.append("--keep-markup")
    if accept_all:
        cmd.append("--accept-all")
    if author:
        cmd.extend(["--author", author])
    try:
        r = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=ADEU_TIMEOUT,
        )
        return {
            "status": "ok" if r.returncode == 0 and dest.is_file() else "error",
            "engine": "adeu-cli",
            "source": str(src),
            "output": str(dest) if dest.is_file() else None,
            "returncode": r.returncode,
            "stderr": (r.stderr or "")[:500],
            "hitl_required": True,
            "elapsed_ms": round((time.perf_counter() - t0) * 1000, 1),
        }
    except Exception as exc:  # noqa: BLE001
        return {"status": "error", "error": str(exc)[:400]}


def markdown_to_docx(
    markdown: str,
    *,
    out_path: str | Path,
    title: str = "A.I.D.A. draft",
) -> dict[str, Any]:
    """Create a simple DOCX from Markdown/plain text (python-docx) for adeu redline."""
    try:
        from docx import Document
    except ImportError:
        return {
            "status": "error",
            "error": "python-docx not installed",
            "hint": "pip install python-docx",
        }
    dest = Path(out_path).expanduser().resolve()
    dest.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    if title:
        doc.add_heading(title[:200], level=1)
    # Very light MD: headings + paragraphs
    for line in (markdown or "").splitlines():
        s = line.rstrip()
        if not s:
            continue
        if s.startswith("### "):
            doc.add_heading(s[4:].strip(), level=3)
        elif s.startswith("## "):
            doc.add_heading(s[3:].strip(), level=2)
        elif s.startswith("# "):
            doc.add_heading(s[2:].strip(), level=1)
        elif s.startswith("- ") or s.startswith("* "):
            doc.add_paragraph(s[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(s)
    doc.save(str(dest))
    return {
        "status": "ok",
        "engine": "python-docx",
        "output": str(dest),
        "bytes": dest.stat().st_size,
        "note": "Scaffold DOCX for adeu redline; not a full Markdown converter",
    }


def from_brief(
    brief_md_path: str | Path | None = None,
    *,
    markdown: str | None = None,
    out_dir: str | Path,
    stem: str = "brief",
    title: str = "Care / advocacy draft",
    edits: list[dict[str, Any]] | None = None,
    author: str | None = None,
) -> dict[str, Any]:
    """Brief MD → DOCX scaffold → optional adeu redline apply."""
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    if brief_md_path and not markdown:
        p = Path(brief_md_path).expanduser().resolve()
        if not p.is_file():
            return {"status": "error", "error": f"brief not found: {p}"}
        markdown = p.read_text(encoding="utf-8")
    if not (markdown or "").strip():
        return {"status": "error", "error": "no markdown content"}

    draft_path = out_dir / f"{stem}.draft.docx"
    created = markdown_to_docx(markdown or "", out_path=draft_path, title=title)
    if created.get("status") != "ok":
        return created

    result: dict[str, Any] = {
        "status": "ok",
        "draft_docx": created.get("output"),
        "create": created,
        "hitl_required": True,
        "decision_authority": "prepare_only",
    }
    if edits:
        redlined = out_dir / f"{stem}.redlined.docx"
        applied = apply_edits(
            draft_path,
            edits,
            out_path=redlined,
            author=author,
        )
        result["redline"] = applied
        result["redlined_docx"] = applied.get("output")
        if applied.get("status") != "ok":
            result["status"] = "partial"
    return result
